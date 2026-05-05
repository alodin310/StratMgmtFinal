from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, size, bold=False, color=None):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, 20, bold=True, color=(31, 73, 125))  # dark blue
    p.paragraph_format.space_after = Pt(4)

def add_chapter_heading(text):
    doc.add_paragraph()
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_font(run, 15, bold=True, color=(17, 85, 161))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)

def add_section_heading(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        set_font(run, 12, bold=True, color=(0, 0, 0))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

def add_row(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25)
    label_run = p.add_run(f"{label}: ")
    set_font(label_run, 11, bold=True)
    val_run = p.add_run(value)
    set_font(val_run, 11)

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════════════

add_title("BCOR 4970 — Strategic Management")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Final Exam Study Guide  •  Chapters 5, 6, 8, 10, 11")
set_font(run, 12, color=(89, 89, 89))

add_divider()

# ── CHAPTER 5 ─────────────────────────────────────────────────────────────────
add_chapter_heading("Chapter 5: Shared Value and Competitive Advantage")

add_section_heading("Shared Value")
add_row("Definition",             "Creating economic value while simultaneously solving social problems")
add_row("Mechanism",              "Align firm activities with the needs of a broad set of stakeholders")
add_row("Competitive Advantage",  "Stronger reputation and long-term customer demand")
add_row("Example",                "Patagonia's sustainability strategy")

add_section_heading("Value Creation Framework")
add_row("Definition",             "Value = Willingness to Pay (WTP) − Cost")
add_row("Mechanism",              "Increase WTP or reduce cost to widen the value gap over competitors")
add_row("Competitive Advantage",  "Larger value gap than rivals")
add_row("Example",                "Apple vs. generic electronics")

add_section_heading("Balanced Scorecard")
add_row("Definition",             "Performance measurement tool integrating financial and non-financial metrics")
add_row("Mechanism",              "Tracks four perspectives: Financial, Customer, Internal Processes, Learning & Growth")
add_row("Competitive Advantage",  "Aligns day-to-day execution with long-term strategy")
add_row("Example",                "Tracking customer satisfaction alongside profit margins")

add_divider()

# ── CHAPTER 6 ─────────────────────────────────────────────────────────────────
add_chapter_heading("Chapter 6: Business Strategy")

add_section_heading("Cost Leadership")
add_row("Definition",             "Producing acceptable value at the lowest cost in the industry")
add_row("Mechanism",              "Economies of scale, efficient operations, tight cost controls")
add_row("Competitive Advantage",  "Ability to survive and win price wars; higher margins at market price")
add_row("Example",                "Aldi, Walmart")

add_section_heading("Differentiation")
add_row("Definition",             "Offering unique value that customers are willing to pay a premium for")
add_row("Mechanism",              "Innovation, superior branding, quality, or service — NOT lowest cost")
add_row("Competitive Advantage",  "Higher willingness to pay → premium pricing and higher margins")
add_row("Example",                "Nike, Starbucks, luxury brands")

add_section_heading("Blue Ocean Strategy")
add_row("Definition",             "Pursuing differentiation and low cost simultaneously to create new market space")
add_row("Mechanism",              "Redefine industry value drivers; eliminate, reduce, raise, create attributes")
add_row("Competitive Advantage",  "Uncontested market space with no direct competitors")
add_row("Example",                "JetBlue — low cost + quality experience")

add_divider()

# ── CHAPTER 8 ─────────────────────────────────────────────────────────────────
add_chapter_heading("Chapter 8: Corporate Strategy")

add_section_heading("Corporate Strategy (Overview)")
add_row("Definition",             "Decisions about which industries and markets the firm competes in")
add_row("Mechanism",              "Allocate resources across business units to maximize overall performance")
add_row("Competitive Advantage",  "Growth and synergy across the portfolio")
add_row("Example",                "Amazon's expansion across retail, cloud, logistics, and media")

add_section_heading("Vertical Integration")
add_row("Definition",             "Owning multiple stages of the value chain (upstream suppliers or downstream distribution)")
add_row("Mechanism",              "Internalize activities previously done by outside firms")
add_row("Competitive Advantage",  "Cost control, quality assurance, and proprietary knowledge")
add_row("Example",                "Tesla (owns manufacturing, sales, charging network)")

add_section_heading("Outsourcing")
add_row("Definition",             "Reducing vertical integration by relying on external suppliers for non-core activities")
add_row("Mechanism",              "Contract out activities outside core competencies")
add_row("Competitive Advantage",  "Focus resources on core activities; lower overhead")
add_row("Example",                "Nike outsources manufacturing, focuses on design and marketing")

add_section_heading("Diversification")
add_row("Definition",             "Entering new industries or markets beyond the firm's current business")
add_row("Mechanism",              "Leverage existing competencies (related) or build new ones (unrelated)")
add_row("Competitive Advantage",  "Risk reduction and/or synergy — but hardest to execute (new market + new capabilities)")
add_row("Example",                "Amazon AWS (related); tech firm entering pharma (unrelated — very difficult)")

add_section_heading("Economies of Scale")
add_row("Definition",             "Cost advantages gained by increasing output volume")
add_row("Mechanism",              "Spread fixed costs across more units → lower cost per unit")
add_row("Competitive Advantage",  "Lower unit cost → pricing power or higher margins")
add_row("Example",                "Walmart")

add_divider()

# ── CHAPTER 10 ────────────────────────────────────────────────────────────────
add_chapter_heading("Chapter 10: Global Strategy")

add_section_heading("Global Strategy (Overview)")
add_row("Definition",             "How a firm achieves and sustains competitive advantage across international markets")
add_row("Competitive Advantage",  "Access to global efficiencies, new markets, and resources")
add_row("Example",                "IKEA's standardized global model")

add_section_heading("CAGE Framework")
add_row("Definition",             "Framework measuring \"distance\" between countries along four dimensions")
add_row("Mechanism",              "Evaluates Cultural, Administrative, Geographic, and Economic barriers to entry")
add_row("Competitive Advantage",  "Better international location and market-entry decisions")
add_row("Example",                "US expansion to Canada — low CAGE distance; US to China — high CAGE distance")

add_section_heading("Cost–Responsiveness Framework (Integration-Responsiveness)")
add_row("Definition",             "Trade-off between global cost reduction and local market responsiveness")
add_row("Mechanism",              "Determines which international strategy type to use (global, multidomestic, transnational)")
add_row("Competitive Advantage",  "Alignment between strategy type and market demands")
add_row("Example",                "McDonald's local menus = multidomestic (high local responsiveness)")

add_section_heading("Multidomestic Strategy")
add_row("Definition",             "High local responsiveness — adapt products and operations to each local market")
add_row("Mechanism",              "Decentralize decisions; tailor offerings by country")
add_row("Competitive Advantage",  "Better customer fit in diverse markets")
add_row("Example",                "McDonald's (local menu items by country)")

add_divider()

# ── CHAPTER 11 ────────────────────────────────────────────────────────────────
add_chapter_heading("Chapter 11: Organizational Design — Structure, Culture & Control")

add_section_heading("Organizational Design (Overview)")
add_row("Definition",             "How a firm structures activities, culture, and controls to execute strategy")
add_row("Mechanism",              "Coordinates activities and aligns incentives across the organization")
add_row("Competitive Advantage",  "Execution efficiency — strategy only works if the org is designed to deliver it")
add_row("Example",                "Amazon's decentralized \"two-pizza team\" structure")

add_section_heading("Organizational Culture")
add_row("Definition",             "Shared values, norms, and behaviors within a firm")
add_row("Mechanism",              "Creates complex, informal coordination that is hard for competitors to imitate")
add_row("Competitive Advantage",  "Sustained competitive advantage — culture is difficult to copy")
add_row("Example",                "Southwest Airlines' culture of employee empowerment and fun")

add_section_heading("Centralization vs. Decentralization")
add_row("Centralization",         "Decisions concentrated at the top — slower, less responsive, but more control")
add_row("Decentralization",       "Decisions pushed down — faster, more responsive, better local adaptation")
add_row("Key Risk",               "Over-centralization leads to inertia and slowed decision-making (e.g., large bureaucracies)")

add_section_heading("Inertia")
add_row("Definition",             "Organizational resistance to change — tendency to maintain status quo")
add_row("Mechanism",              "Embedded routines, culture, and structures prevent adaptation")
add_row("Competitive Advantage",  "Firms that overcome inertia can adapt; those that don't lose advantage over time")
add_row("Example",                "Kodak — failed to adapt to digital photography despite inventing it")

add_section_heading("Business Model")
add_row("Definition",             "How a firm translates its strategy into economic value — defines value creation, delivery, and capture")
add_row("Mechanism",              "Aligns resources, activities, and revenue model to execute strategy")
add_row("Competitive Advantage",  "Ensures organizational design supports strategic goals")
add_row("Example",                "Amazon's flywheel — low prices → traffic → sellers → more selection → lower costs")

add_divider()

# ── EXAM FOCUS ────────────────────────────────────────────────────────────────
add_chapter_heading("Exam Focus — Key Distinctions & Common Traps")

add_section_heading("Differentiation vs. Cost Leadership")
add_row("Trap",      "Differentiation does NOT mean lowest cost — it means highest perceived value")
add_row("Cost Leader", "Lowest cost structure with acceptable value → wins price wars (e.g., Aldi)")
add_row("Differentiator", "Premium pricing through superior value → higher margins (e.g., Starbucks, Luxury brands)")

add_section_heading("Diversification Difficulty Spectrum")
add_row("Easiest",   "Same market, same competencies (incremental growth)")
add_row("Harder",    "Related diversification — new market OR new competencies")
add_row("Hardest",   "Unrelated diversification — new market AND new competencies (e.g., tech firm entering pharma)")

add_section_heading("Multiple Strategies Can Win in the Same Industry")
add_row("Key Point", "Different strategic positions can both succeed — there is not one winner per industry")
add_row("Example",   "Aldi (cost leader) and Whole Foods (differentiator) both succeed in grocery retail")

add_section_heading("Commoditization Threatens Differentiation")
add_row("Definition", "When a formerly differentiated product becomes indistinguishable from competitors'")
add_row("Effect",     "Reduces uniqueness → erodes premium pricing → margins compress")
add_row("Example",    "Smartphones — once highly differentiated, now largely commoditized")

add_divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Good luck on May 7!")
set_font(run, 11, bold=True, color=(89, 89, 89))

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/anderslodin/isba-4715/StratMgmtFinal/docs/StratMgmt_Final_StudyGuide.docx"
doc.save(out)
print(f"Saved: {out}")
