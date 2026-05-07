from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

def heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def question(num, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {text}")
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p

def choice(letter, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(f"{letter})  {text}")
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(1)
    return p

def blank_lines(n=4):
    for _ in range(n):
        doc.add_paragraph()

def page_break():
    doc.add_page_break()

# ── Title ────────────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("BCOR 4970 — Strategic Management")
r.bold = True
r.font.size = Pt(16)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run("Practice Final Exam  •  Chapters 5, 6, 8, 10, 11")
r2.font.size = Pt(12)

doc.add_paragraph()

instructions = doc.add_paragraph()
instructions.add_run(
    "Instructions: This exam contains 20 multiple choice questions and 5 short answer questions. "
    "For multiple choice, circle or highlight the letter of the best answer. "
    "For short answer, write your response in the space provided. "
    "An answer key is located after a page break at the end of this document."
)
instructions.runs[0].font.size = Pt(10)
instructions.runs[0].italic = True

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION I — MULTIPLE CHOICE
# ════════════════════════════════════════════════════════════════════════════

heading("SECTION I — Multiple Choice  (20 Questions)")

# ── Ch 5 ────────────────────────────────────────────────────────────────────

heading("Chapter 5: Shared Value and Competitive Advantage", level=2)

question(1, "A firm produces a product at a cost of $35. A customer's maximum willingness "
            "to pay is $80, and the firm charges $60. What is the TOTAL value created by "
            "this transaction?")
choice("A", "$25")
choice("B", "$45")
choice("C", "$60")
choice("D", "$80")

question(2, "Using the same scenario (cost $35, WTP $80, price $60), how much value does "
            "the CUSTOMER capture?")
choice("A", "$45")
choice("B", "$35")
choice("C", "$25")
choice("D", "$20")

question(3, "Patagonia's commitment to environmental sustainability has strengthened its "
            "brand, increased customer loyalty, and driven long-term revenue growth. "
            "This is BEST described as:")
choice("A", "Corporate social responsibility — a cost the firm absorbs to improve public image")
choice("B", "Shared value — social and economic value creation are mutually reinforcing")
choice("C", "Differentiation strategy — premium pricing through sustainability branding")
choice("D", "Blue Ocean strategy — creating an uncontested market in ethical apparel")

question(4, "A company's strategy scorecard tracks: (1) return on equity, (2) customer "
            "satisfaction scores, (3) on-time delivery rates, and (4) employee training "
            "completion. Which framework does this BEST represent?")
choice("A", "CAGE Framework")
choice("B", "Integration-Responsiveness Framework")
choice("C", "Balanced Scorecard")
choice("D", "Value Chain Analysis")

# ── Ch 6 ────────────────────────────────────────────────────────────────────

heading("Chapter 6: Business Strategy", level=2)

question(5, "Which of the following is TRUE about cost leadership strategy?")
choice("A", "Cost leaders must charge the lowest prices in the industry")
choice("B", "Cost leaders achieve competitive advantage through the lowest cost structure, "
            "not necessarily the lowest price")
choice("C", "Cost leaders differentiate by offering the most value to customers")
choice("D", "Cost leadership requires a higher WTP than competitors")

question(6, "A luxury car brand spends heavily on design, engineering, and customer "
            "experience to justify a price premium. A value brand produces basic, "
            "reliable vehicles at the lowest cost per unit. Which statement BEST "
            "compares these firms?")
choice("A", "The luxury brand is a cost leader because its cars are more efficient to produce")
choice("B", "Both firms are pursuing differentiation but in different market segments")
choice("C", "The luxury brand raises WTP; the value brand reduces cost — both widen the "
            "value gap differently")
choice("D", "The value brand is pursuing Blue Ocean strategy by avoiding direct competition")

question(7, "JetBlue launched with low fares AND quality amenities (leather seats, "
            "live TV, more legroom) in an industry where competitors chose either low "
            "cost OR comfort. This is BEST described as:")
choice("A", "Cost leadership — fares were below legacy carriers")
choice("B", "Differentiation — the amenities justified a premium price")
choice("C", "Blue Ocean strategy — combining low cost and differentiation to create new "
            "market space")
choice("D", "Multidomestic strategy — tailoring the experience to US travelers")

question(8, "Over the past decade, smartphones have become increasingly similar in "
            "features, performance, and price across major brands. Premium pricing has "
            "eroded. This trend BEST illustrates:")
choice("A", "Vertical integration")
choice("B", "Commoditization")
choice("C", "Organizational inertia")
choice("D", "Economies of scale")

question(9, "Whole Foods (premium organic grocery) and Aldi (no-frills discount grocery) "
            "both operate profitably in the same industry. What does this BEST illustrate?")
choice("A", "Differentiation always outperforms cost leadership over the long run")
choice("B", "Industry structure dictates which single strategy must be used to compete")
choice("C", "Multiple distinct strategic positions can coexist successfully in the same industry")
choice("D", "Blue Ocean strategy eliminates competition between firms")

question(10, "A differentiating firm launches a high-end product but fails to manage "
             "production costs, resulting in margins no better than competitors. "
             "What went wrong strategically?")
choice("A", "The firm incorrectly pursued cost leadership instead of differentiation")
choice("B", "The firm raised WTP but failed to maintain sufficient cost discipline — "
            "differentiators still manage cost, they just don't lead with it")
choice("C", "The firm entered a Blue Ocean market without sufficient resources")
choice("D", "The firm over-invested in shared value at the expense of profit")

# ── Ch 8 ────────────────────────────────────────────────────────────────────

heading("Chapter 8: Corporate Strategy", level=2)

question(11, "Which BEST distinguishes corporate strategy from business strategy?")
choice("A", "Corporate strategy determines HOW to compete; business strategy determines "
            "WHERE to compete")
choice("B", "Corporate strategy determines WHERE to compete; business strategy determines "
            "HOW to compete")
choice("C", "Corporate strategy applies to large firms; business strategy applies to startups")
choice("D", "They are interchangeable terms describing the same set of decisions")

question(12, "Tesla owns its manufacturing plants, retail showrooms, and charging network. "
             "Nike outsources all manufacturing and focuses on design, marketing, and brand. "
             "Which statement BEST explains this strategic difference?")
choice("A", "Tesla is more profitable than Nike because vertical integration always reduces cost")
choice("B", "Nike's strategy is inferior because outsourcing cedes control over quality")
choice("C", "Tesla integrates to gain control and quality assurance; Nike outsources to focus "
            "on core competencies and gain flexibility")
choice("D", "Both firms are pursuing the same corporate strategy in different industries")

question(13, "A software company acquires a direct competitor (same software, same market). "
             "Later, it attempts to enter the pharmaceutical industry. How do these two "
             "diversification moves compare in difficulty?")
choice("A", "They are equally difficult because diversification always requires new investment")
choice("B", "The pharma entry is harder — it requires both new competencies AND a new market")
choice("C", "The acquisition is harder because integrating a competitor requires more resources")
choice("D", "Unrelated diversification is always easier because there is no direct competition")

question(14, "Walmart can negotiate dramatically lower per-unit costs from suppliers than "
             "regional grocery chains. This advantage is primarily driven by:")
choice("A", "Multidomestic strategy")
choice("B", "Vertical integration of the supply chain")
choice("C", "Economies of scale — spreading fixed costs across massive volume")
choice("D", "Related diversification into wholesale distribution")

question(15, "A firm outsources its customer service operations to a third-party provider "
             "so internal teams can focus entirely on product innovation. This decision "
             "is BEST described as:")
choice("A", "Unrelated diversification into the service sector")
choice("B", "Reducing vertical integration to focus resources on core competencies")
choice("C", "A Blue Ocean move to differentiate through outsourced expertise")
choice("D", "Decentralizing organizational structure to improve responsiveness")

# ── Ch 10 ────────────────────────────────────────────────────────────────────

heading("Chapter 10: Global Strategy", level=2)

question(16, "A US company is evaluating expansion into Canada versus Vietnam. "
             "According to the CAGE framework, Canada would present LOWER distance "
             "primarily because of:")
choice("A", "Geographic proximity alone")
choice("B", "Shared language, legal traditions, colonial ties, and similar economic "
            "development — lower distance across all four CAGE dimensions")
choice("C", "Lower administrative distance due to NAFTA/USMCA trade agreements only")
choice("D", "Canada and Vietnam present equal CAGE distance from the US")

question(17, "IKEA sells the same furniture designs, uses identical store layouts, and "
             "maintains the same brand experience across 50+ countries. This is an "
             "example of:")
choice("A", "Multidomestic strategy")
choice("B", "Transnational strategy")
choice("C", "Global strategy")
choice("D", "International strategy")

question(18, "McDonald's offers a McSpicy Paneer in India, a Teriyaki Burger in Japan, "
             "and a McRib in the US — while maintaining global brand standards. "
             "Which strategy type BEST describes this approach?")
choice("A", "Global strategy — leveraging brand standardization for global efficiency")
choice("B", "Multidomestic strategy — adapting products significantly to local markets")
choice("C", "Transnational strategy — achieving high global integration and high local responsiveness")
choice("D", "Blue Ocean strategy — creating new market space in each country")

question(19, "A firm pursuing transnational strategy faces greater execution challenges "
             "than one pursuing a pure global or multidomestic strategy because it must:")
choice("A", "Standardize all products globally while entering new markets")
choice("B", "Adapt completely to every local market with no global coordination")
choice("C", "Simultaneously achieve high global integration AND high local responsiveness — "
            "two objectives that naturally conflict")
choice("D", "Minimize CAGE distance by only entering culturally similar markets")

# ── Ch 11 ────────────────────────────────────────────────────────────────────

heading("Chapter 11: Organizational Design", level=2)

question(20, "Kodak invented digital photography technology in 1975 but failed to "
             "commercialize it and went bankrupt as digital cameras took over the "
             "market. This failure is BEST explained by:")
choice("A", "Lack of financial resources to develop digital technology")
choice("B", "Poor individual leadership decisions with no structural basis")
choice("C", "Organizational inertia — embedded processes, culture, and incentives aligned "
            "to film prevented the firm from adapting even when it saw the threat")
choice("D", "Blue Ocean competitors creating uncontested digital market space")

# ════════════════════════════════════════════════════════════════════════════
# SECTION II — SHORT ANSWER
# ════════════════════════════════════════════════════════════════════════════

page_break()
heading("SECTION II — Short Answer  (5 Questions)")

p = doc.add_paragraph()
p.add_run(
    "For each question, write your response in the space provided. "
    "Strong answers include: a clear definition, the mechanism (how/why it works), "
    "and at least one real-world example."
).italic = True
p.runs[0].font.size = Pt(10)

doc.add_paragraph()

# SA 1
heading("Question 21", level=2)
question(21, "Define the value equation (Value = WTP − Cost). "
             "Explain how a cost leader and a differentiator each use this equation to "
             "build competitive advantage — and what happens to the equation when a "
             "formerly differentiated product becomes commoditized. "
             "Use a real-world example for each strategy.")
blank_lines(8)

# SA 2
heading("Question 22", level=2)
question(22, "Two meal-kit companies compete in the same market. "
             "Company A offers premium, chef-curated recipes with organic ingredients at $14/meal. "
             "Company B offers simple, reliable recipes with standard ingredients at $6/meal. \n\n"
             "(a) Identify the business strategy each company is pursuing and justify your answer.\n"
             "(b) Company A's ingredients and packaging become widely available through grocery "
             "stores. How does this threaten Company A's strategy, and what should it do?")
blank_lines(10)

# SA 3
heading("Question 23", level=2)
question(23, "Describe the spectrum from full vertical integration to full outsourcing. "
             "Under what conditions should a firm choose to vertically integrate, and "
             "when does outsourcing create MORE competitive advantage? "
             "Use Tesla and Nike as your examples and explain why each firm's choice fits its strategy.")
blank_lines(8)

# SA 4
heading("Question 24", level=2)
question(24, "A US athletic wear brand (similar to Nike) is evaluating international "
             "expansion into Germany and India.\n\n"
             "(a) Apply the CAGE framework to compare these two markets. Identify which "
             "CAGE dimensions create the greatest distance for each country.\n"
             "(b) Recommend which global strategy type (global, multidomestic, "
             "transnational, or international) the brand should use, and justify your "
             "recommendation using the Integration-Responsiveness framework.")
blank_lines(10)

# SA 5
heading("Question 25", level=2)
question(25, "Explain why organizational culture is considered a source of SUSTAINED "
             "competitive advantage — not just a temporary one. "
             "How does inertia relate to culture, and what does the Kodak example reveal "
             "about the relationship between past success, culture, and the inability to adapt? "
             "Use Southwest Airlines as a contrasting example.")
blank_lines(8)

# ════════════════════════════════════════════════════════════════════════════
# ANSWER KEY
# ════════════════════════════════════════════════════════════════════════════

page_break()

ak_title = doc.add_paragraph()
ak_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ak_title.add_run("— ANSWER KEY —")
r.bold = True
r.font.size = Pt(14)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.add_run("Do not look at this until you have completed the exam.").italic = True
note.runs[0].font.size = Pt(10)

doc.add_paragraph()

heading("Multiple Choice Answers", level=2)

mc_answers = [
    ("1",  "B — Value = WTP − Cost = $80 − $35 = $45"),
    ("2",  "D — Consumer surplus = WTP − Price = $80 − $60 = $20"),
    ("3",  "B — Shared value: social and economic goals are mutually reinforcing, not a trade-off"),
    ("4",  "C — Balanced Scorecard tracks financial, customer, internal process, and learning metrics"),
    ("5",  "B — Cost leaders have the lowest cost STRUCTURE; price is a strategic choice, not a requirement"),
    ("6",  "C — Each firm widens the value gap differently: one via WTP, one via cost"),
    ("7",  "C — Blue Ocean: simultaneous low cost + differentiation creates uncontested market space"),
    ("8",  "B — Commoditization: formerly differentiated products become indistinguishable"),
    ("9",  "C — Multiple strategic positions (cost leader + differentiator) can both succeed in one industry"),
    ("10", "B — Differentiators still manage cost; raising WTP without cost discipline destroys margins"),
    ("11", "B — Corporate = WHERE to compete. Business = HOW to compete"),
    ("12", "C — Tesla integrates for control/quality. Nike outsources for flexibility/focus on core"),
    ("13", "B — Unrelated diversification (new market + new competencies) is always harder"),
    ("14", "C — Walmart's volume spreads fixed cost → lower cost per unit = economies of scale"),
    ("15", "B — Reducing vertical integration by outsourcing non-core activity to focus on core competency"),
    ("16", "B — Canada scores lower on all four CAGE dimensions, not just geography"),
    ("17", "C — Global strategy: standardize across all markets for global efficiency"),
    ("18", "B — Multidomestic: local product adaptation despite global brand"),
    ("19", "C — Transnational must achieve both high integration AND high responsiveness simultaneously"),
    ("20", "C — Organizational inertia: success in film embedded routines that prevented digital adaptation"),
]

for num, ans in mc_answers:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"Q{num}:  ").bold = True
    p.add_run(ans).font.size = Pt(10)

doc.add_paragraph()
heading("Short Answer — Key Points", level=2)

sa_answers = [
    ("21 — Value Equation + Commoditization",
     [
         "Value = WTP − Cost is the total economic value created; it is split between customer surplus (WTP − Price) and firm profit (Price − Cost).",
         "Cost leader (e.g., Aldi): reduces cost while maintaining acceptable WTP → widens gap from the bottom.",
         "Differentiator (e.g., Apple): raises WTP through innovation/branding while managing cost → widens gap from the top.",
         "Commoditization collapses WTP toward competitors' level, eroding the differentiator's premium. "
         "If cost doesn't fall simultaneously, the value gap shrinks and competitive advantage disappears.",
     ]),
    ("22 — Business Strategy + Commoditization Threat",
     [
         "(a) Company A = Differentiation — high WTP through premium ingredients/experience → justifies $14 price. "
         "Company B = Cost Leadership — acceptable value at lowest cost → $6 price wins price-sensitive buyers.",
         "(b) Grocery availability commoditizes Company A's ingredients → WTP falls toward Company B's level. "
         "Company A must re-differentiate (exclusive recipes, chef brand, community/subscription experience) "
         "or risk margin collapse. Cannot compete on cost — that's Company B's game.",
     ]),
    ("23 — Vertical Integration vs. Outsourcing",
     [
         "Spectrum: full integration (own all stages) ↔ full outsourcing (contract everything but core).",
         "Integrate when: activity is core to competitive advantage, quality/cost control is critical, "
         "proprietary knowledge must be protected. → Tesla: manufacturing precision and the charging network "
         "are core to its differentiated product; owning them ensures quality and locks out competitors.",
         "Outsource when: activity is non-core, external suppliers do it better/cheaper, "
         "flexibility and scale are more valuable than control. → Nike: outsources manufacturing to focus "
         "resources on design and brand — its actual source of competitive advantage.",
     ]),
    ("24 — CAGE + IR Framework",
     [
         "(a) Germany: lower cultural distance (developed economy, Western consumer norms), "
         "lower economic distance (similar income/infrastructure), moderate administrative distance. "
         "India: higher cultural distance (language, religion, norms), higher economic distance "
         "(income levels, infrastructure), higher administrative distance (regulations, IP environment). "
         "India = higher overall CAGE distance.",
         "(b) For a global brand like Nike: Germany → closer to global/international strategy "
         "(less adaptation needed). India → higher local responsiveness required (price sensitivity, "
         "cultural fit) → multidomestic elements. Overall recommendation: transnational — "
         "maintain global brand integration while adapting product lines and pricing to local markets. "
         "Justification: high global integration (brand consistency) + moderate-to-high local "
         "responsiveness (pricing, product range) → upper-right quadrant of IR framework.",
     ]),
    ("25 — Culture as Sustained Advantage + Inertia",
     [
         "Culture = shared values, norms, and behaviors embedded over years. It is hard to imitate "
         "because it cannot be purchased, copied from a manual, or transferred — it must be grown. "
         "This makes it a source of SUSTAINED (not just temporary) competitive advantage.",
         "Inertia = organizational resistance to change. Culture contributes to inertia: "
         "the same deeply embedded norms that drive performance in a stable environment "
         "prevent adaptation when the environment changes.",
         "Kodak: success in film created strong cultural and structural alignment to the film business. "
         "When digital disrupted the market, that alignment became a trap. Kodak had the technology "
         "but not the organizational will or structure to deploy it.",
         "Southwest Airlines (contrast): culture of employee empowerment and service is deliberately "
         "built and reinforced — it drives performance AND is nearly impossible for competitors to replicate "
         "even if they know exactly what Southwest does.",
     ]),
]

for title_str, bullets in sa_answers:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(f"Q{title_str}")
    r.bold = True
    r.font.size = Pt(11)
    for b in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.add_run(b).font.size = Pt(10)
        bp.paragraph_format.space_after = Pt(2)

# ── Save ─────────────────────────────────────────────────────────────────────

out_path = "StratMgmt_PracticeExam.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
