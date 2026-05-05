from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── helpers ───────────────────────────────────────────────────────────────────
def font(run, size, bold=False, italic=False, color=None):
    run.font.name   = "Calibri"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    font(r, 20, bold=True, color=(31, 73, 125))
    p.paragraph_format.space_after = Pt(2)

def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    font(r, 11, italic=True, color=(89, 89, 89))
    p.paragraph_format.space_after = Pt(8)

def ch_heading(text):
    doc.add_paragraph()
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        font(r, 14, bold=True, color=(17, 85, 161))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def sub(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        font(r, 11, bold=True, color=(0, 0, 0))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)

def bullet(text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    font(r, 10.5)

def bullet_bold(label, rest, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(label)
    font(r1, 10.5, bold=True)
    r2 = p.add_run(rest)
    font(r2, 10.5)

def callout(label, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{label}  ")
    font(r1, 10.5, bold=True, color=color)
    r2 = p.add_run(text)
    font(r2, 10.5, italic=True)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "BBBBBB")
    pBdr.append(bot)
    pPr.append(pBdr)

def plain(text, size=10.5, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    font(r, size, bold=bold, italic=italic)

# ══════════════════════════════════════════════════════════════════════════════
# CONTENT
# ══════════════════════════════════════════════════════════════════════════════

title("BCOR 4970 — Strategic Management")
subtitle("Final Exam Strategy Guide  •  3 Days  •  4 Questions  •  Chapters 5, 6, 8, 10, 11")

divider()

# ── THE BIG PICTURE ───────────────────────────────────────────────────────────
ch_heading("The Overarching Theme of This Course")

plain(
    "Every single chapter answers one question: How do firms gain and sustain competitive advantage? "
    "Learn to see each chapter as one layer of that answer — they build on each other.",
    bold=True
)

bullet_bold("Ch 5 — ", "WHAT is competitive advantage? (Value = WTP − Cost)")
bullet_bold("Ch 6 — ", "HOW do you compete? (Cost leader vs. Differentiator vs. Blue Ocean)")
bullet_bold("Ch 8 — ", "WHERE do you compete? (Which industries, how integrated, how diversified?)")
bullet_bold("Ch 10 — ", "WHERE in the world? (Which countries, which global strategy?)")
bullet_bold("Ch 11 — ", "HOW do you execute? (Structure, culture, controls to deliver the strategy)")

callout("KEY INSIGHT:", "Strategy without execution is just a plan. Ch 11 is WHY strategy fails or succeeds in practice.", (31, 73, 125))

divider()

# ── STUDY PLAN ────────────────────────────────────────────────────────────────
ch_heading("3-Day Study Plan")

sub("Day 1 — Understand the Frameworks (Ch 5 + Ch 6)")
bullet("Ch 5 first — it sets the vocabulary (WTP, cost, value gap, competitive advantage) that every other chapter uses")
bullet("Ch 6 second — the most testable chapter; know cost leadership vs. differentiation cold")
bullet("Goal: be able to explain each concept in 1–2 sentences without notes")

sub("Day 2 — Apply the Frameworks (Ch 8 + Ch 10)")
bullet("Ch 8 — focus on the spectrum: full vertical integration → outsourcing, and related vs. unrelated diversification")
bullet("Ch 10 — memorize CAGE + the 4 international strategy types and when to use each")
bullet("Goal: practice applying frameworks to mini scenarios (e.g., 'Should Company X enter Country Y?')")

sub("Day 3 — Connect and Review (Ch 11 + Exam Prep)")
bullet("Ch 11 — understand why org design is the final piece that makes strategy real")
bullet("Review the Exam Focus section of your study guide — drill the common traps")
bullet("Goal: write out 2–3 sentence answers for likely essay questions without looking at notes")

divider()

# ── CHAPTER 5 ─────────────────────────────────────────────────────────────────
ch_heading("Chapter 5 — Shared Value & Competitive Advantage")

sub("What Actually Matters")
bullet("The value equation:  Value = WTP − Cost  — this is the foundation of the entire course")
bullet("Competitive advantage = your value gap is bigger than competitors'")
bullet("Shared value = you can do social good AND make money — they reinforce each other")
bullet("Balanced Scorecard = strategy needs more than financial metrics to stay on track")

sub("How to Study It")
bullet("Memorize the value equation. Draw it as a simple diagram: WTP on top, Cost on bottom, the gap in between")
bullet("For shared value: think Patagonia — sustainability builds reputation → demand → revenue. Social + economic = aligned")
bullet("For Balanced Scorecard: remember the 4 boxes — Financial, Customer, Internal, Learning. They all connect to strategy")

sub("Memory Tricks")
bullet_bold("Value equation: ", "'WTP minus C equals V' — Willingness To Pay minus Cost equals Value gap")
bullet_bold("Balanced Scorecard: ", "FCIL — Financial, Customer, Internal, Learning")
bullet_bold("Shared Value: ", "Doing good = doing well. Not charity, not tradeoff — alignment")

callout("CONNECTS TO:", "Ch 6 (how to widen the value gap via cost or differentiation) and Ch 11 (execution determines if you actually capture that value)", (0, 128, 0))

divider()

# ── CHAPTER 6 ─────────────────────────────────────────────────────────────────
ch_heading("Chapter 6 — Business Strategy")

sub("What Actually Matters")
bullet("The #1 most testable concept: Cost Leadership vs. Differentiation — know the difference cold")
bullet("Cost leader = lowest cost, acceptable value. Wins price wars. Does NOT mean cheapest product")
bullet("Differentiator = highest perceived value, customers pay premium. Does NOT mean low cost")
bullet("Blue Ocean = escape the cost vs. differentiation trade-off. Redefine the game entirely")
bullet("Both strategies can win in the same industry — different positions, not a zero-sum fight")

sub("How to Study It")
bullet("Draw the 2x2: Low Cost / High Cost on one axis, Low Value / High Value on the other")
bullet("Place firms: Aldi (low cost, acceptable value), Starbucks (high value, premium price), JetBlue (Blue Ocean)")
bullet("Practice the trap: 'Differentiation means ignoring cost' — FALSE. Differentiators still manage cost, they just don't lead with it")

sub("Memory Tricks")
bullet_bold("Cost leader: ", "Survive price WARS → think Aldi, Walmart (WAR-mart)")
bullet_bold("Differentiator: ", "Customers PAY MORE → think Nike, Starbucks, Apple")
bullet_bold("Blue Ocean: ", "Create your own ocean — no sharks (competitors). JetBlue left the red ocean")
bullet_bold("Quick test: ", "Ask 'why do customers buy this?' — Value = differentiator. Price = cost leader")

callout("CONNECTS TO:", "Ch 5 (differentiators raise WTP; cost leaders reduce cost — both widen the value gap) and Ch 8 (business strategy is HOW you compete within the industries Ch 8 says to enter)", (0, 128, 0))

divider()

# ── CHAPTER 8 ─────────────────────────────────────────────────────────────────
ch_heading("Chapter 8 — Corporate Strategy")

sub("What Actually Matters")
bullet("Corporate strategy = WHERE to compete (Ch 6 is HOW). These two are always paired on exams")
bullet("Vertical integration spectrum: fully integrated ↔ fully outsourced. Tesla vs. Nike")
bullet("Diversification difficulty: same market + same competencies (easy) → new market + new competencies (hardest)")
bullet("Economies of scale = spread fixed cost → lower unit cost → competitive advantage")
bullet("Outsourcing = strategic choice to focus on core; not a failure, not always the answer")

sub("How to Study It")
bullet("Draw the diversification matrix: competencies (same/new) on one axis, markets (same/new) on the other")
bullet("Place examples: Amazon retail → AWS (same competencies, new market = related). Tech firm → pharma (new everything = unrelated)")
bullet("For vertical integration: ask 'does owning this activity give us control, quality, or cost advantage?' → Tesla yes. Nike no")

sub("Memory Tricks")
bullet_bold("Corporate vs. Business strategy: ", "Corporate = WHAT game to play. Business = HOW to play it")
bullet_bold("Diversification order: ", "Same/Same → Related (one new) → Unrelated (both new = hardest)")
bullet_bold("Tesla = integrate. Nike = outsource: ", "Tesla needs precision control. Nike needs flexibility and scale")
bullet_bold("Economies of scale: ", "Walmart buys SO much that each unit costs less. Spread the fixed cost wide")

callout("CONNECTS TO:", "Ch 6 (once you pick where to compete, Ch 6 tells you how) and Ch 10 (global strategy = corporate strategy applied internationally)", (0, 128, 0))

divider()

# ── CHAPTER 10 ────────────────────────────────────────────────────────────────
ch_heading("Chapter 10 — Global Strategy")

sub("What Actually Matters")
bullet("CAGE Framework — four types of distance that make international expansion harder")
bullet("Integration-Responsiveness trade-off — the central framework for choosing a global strategy type")
bullet("4 global strategy types: Global, Multidomestic, Transnational, International")
bullet("Key distinction: Global strategy = standardize for efficiency. Multidomestic = adapt for local fit")

sub("CAGE Framework — The 4 Distances")
bullet_bold("C — Cultural: ", "Language, religion, values, norms")
bullet_bold("A — Administrative: ", "Laws, regulations, political relationships, colonial ties")
bullet_bold("G — Geographic: ", "Physical distance, lack of shared borders, time zones")
bullet_bold("E — Economic: ", "Wage differences, income levels, infrastructure quality")
bullet("Higher CAGE distance = harder and riskier to enter. US→Canada = low. US→China = high")

sub("4 Global Strategy Types")
bullet_bold("Global: ", "Low local responsiveness + high global efficiency. One product, everywhere. (IKEA)")
bullet_bold("Multidomestic: ", "High local responsiveness + low global efficiency. Adapt everything. (McDonald's local menus)")
bullet_bold("Transnational: ", "High on both — the hardest to execute. Best of both worlds. (Unilever)")
bullet_bold("International: ", "Low on both — export home product with minimal changes. (Early-stage exporters)")

sub("How to Study It")
bullet("Draw the 2x2 with Global Integration on Y-axis and Local Responsiveness on X-axis. Place all 4 strategies")
bullet("For CAGE: practice with real country pairs. Mexico vs. Japan vs. Canada vs. Germany as US expansion targets")
bullet("For exam questions: if they give a company scenario, ask 'how much do they need to adapt locally?' → that tells you the strategy")

sub("Memory Tricks")
bullet_bold("CAGE: ", "Think of distance as CAGE bars — each one traps your ability to expand easily")
bullet_bold("Multidomestic: ", "McDonald's has a McSpicy Paneer in India and a McRib in the US. Different EVERYTHING")
bullet_bold("Global: ", "IKEA is IKEA in Tokyo, Berlin, and Detroit. Same flat-pack, same meatballs")

callout("CONNECTS TO:", "Ch 8 (global strategy IS corporate strategy — just across borders; vertical integration and diversification decisions apply globally too)", (0, 128, 0))

divider()

# ── CHAPTER 11 ────────────────────────────────────────────────────────────────
ch_heading("Chapter 11 — Organizational Design: Structure, Culture & Control")

sub("What Actually Matters")
bullet("Org design = the bridge between strategy and results. Best strategy fails with wrong structure")
bullet("Culture = the hardest competitive advantage to copy → most sustainable")
bullet("Inertia = the enemy. Firms that can't change lose. Kodak invented digital, still failed")
bullet("Centralization vs. decentralization = speed vs. control trade-off")
bullet("Business model = how the strategy actually creates, delivers, and captures value")

sub("How to Study It")
bullet("Think of Ch 11 as the 'why strategies fail' chapter — connect every concept to execution breakdowns")
bullet("For culture: it's informal, deeply embedded, and takes years to build. That's why it's hard to imitate")
bullet("For inertia: the better a firm was at the OLD game, the harder it is to change. Success breeds inertia")
bullet("For business model: value creation (make it) + value delivery (get it to customers) + value capture (make money from it)")

sub("Memory Tricks")
bullet_bold("Structure follows strategy: ", "You design the org AROUND the strategy, not the other way")
bullet_bold("Culture = Southwest: ", "Fun, empowered employees → better service → loyalty → profit. Hard to copy")
bullet_bold("Inertia = Kodak: ", "They HAD digital. Couldn't change. Film habits = organizational inertia")
bullet_bold("Business model = Amazon flywheel: ", "Low prices → traffic → sellers → selection → lower cost → lower prices")
bullet_bold("Centralization trap: ", "Too much top-down = slow decisions = missed opportunities")

callout("CONNECTS TO:", "Every chapter. Ch 11 is the 'execution layer' for all of Ch 5–10. A firm can have perfect strategy (Ch 6), great global positioning (Ch 10), and smart diversification (Ch 8), and still fail if Ch 11 is broken.", (0, 128, 0))

divider()

# ── EXAM STRATEGY ─────────────────────────────────────────────────────────────
ch_heading("Exam Strategy — 4 Questions, ~3.2 Hours")

sub("How the Exam Likely Works")
bullet("4 questions = probably one per major framework cluster. Expect MCQs, essay, and/or a mini case")
bullet("Mini case = they give you a company scenario and ask you to apply a framework. Practice this format")
bullet("Essay = define + explain mechanism + give example. Three parts every time")

sub("The 5 Frameworks You Must Know Cold")
bullet_bold("1. Value equation (Ch 5): ", "Value = WTP − Cost. Know how each strategy affects it")
bullet_bold("2. Cost vs. Differentiation (Ch 6): ", "The central distinction. Never confuse them")
bullet_bold("3. Diversification matrix (Ch 8): ", "Same/new competencies × same/new markets")
bullet_bold("4. CAGE + IR framework (Ch 10): ", "4 distances + 4 global strategy types on the 2x2")
bullet_bold("5. Structure-Culture-Control (Ch 11): ", "Know why culture = sustained advantage, inertia = failure")

sub("Common Exam Traps — Don't Fall For These")
bullet("Differentiation ≠ low cost. Differentiators CONTROL cost but don't LEAD with it")
bullet("Unrelated diversification ≠ growth strategy. It's the hardest and riskiest move")
bullet("Vertical integration ≠ always good. Sometimes outsourcing creates more competitive advantage")
bullet("Global strategy ≠ multidomestic. Global = standardize. Multidomestic = adapt")
bullet("Culture ≠ just 'company vibe.' It's a strategic asset that creates sustained competitive advantage")
bullet("Inertia ≠ laziness. It's structural — embedded in processes, culture, and incentives")

sub("If You Only Have 1 Hour Left Before the Exam")
bullet("Re-read the Exam Focus section of the study guide")
bullet("Write out the value equation and the 4 global strategy types from memory")
bullet("Think of one real-world example for each of the 5 major concepts")
bullet("Remember: the exam tests application, not just definitions — always answer WHY")

divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("You've got this. Strategy is about making clear choices — so is studying. Good luck May 7.")
font(r, 11, italic=True, color=(89, 89, 89))

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/anderslodin/isba-4715/StratMgmtFinal/docs/StratMgmt_ExamStrategy_Guide.docx"
doc.save(out)
print(f"Saved: {out}")
