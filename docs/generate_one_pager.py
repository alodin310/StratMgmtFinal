from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Tight margins (0.5" all around) ─────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(0.45)
section.bottom_margin = Inches(0.45)
section.left_margin   = Inches(0.5)
section.right_margin  = Inches(0.5)

def set_spacing(para, before=0, after=1):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    para.paragraph_format.line_spacing = Pt(11)

def title_line(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    set_spacing(p, before=0, after=2)

def section_head(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    set_spacing(p, before=4, after=1)

def bullet(label, body, indent=0.15):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    if label:
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(9)
    r2 = p.add_run(body)
    r2.font.size = Pt(9)
    set_spacing(p, before=0, after=1)

def trap(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run(f"✗  {text}")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    set_spacing(p, before=0, after=1)

def divider():
    p = doc.add_paragraph("─" * 110)
    p.runs[0].font.size = Pt(6)
    p.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    set_spacing(p, before=2, after=2)

# ════════════════════════════════════════════════════════════════════════════

title_line("BCOR 4970 Strategic Management — Exam Day One-Pager")

divider()

# ── Ch 5 ────────────────────────────────────────────────────────────────────
section_head("CH 5 — Shared Value & Competitive Advantage")
bullet("Core equation", "Value = WTP − Cost  (total pie created; price just splits it between customer surplus & firm profit)")
bullet("Competitive advantage", "Your value gap (WTP − Cost) is BIGGER than rivals'")
bullet("Shared value", "Social good + economic value reinforce each other — NOT charity, NOT a trade-off  (Patagonia)")
bullet("Balanced Scorecard", "4 perspectives: Financial · Customer · Internal Processes · Learning & Growth")

divider()

# ── Ch 6 ────────────────────────────────────────────────────────────────────
section_head("CH 6 — Business Strategy  (MOST TESTABLE CHAPTER)")
bullet("Cost Leadership", "Lowest cost structure → survive price wars, earn margins at market price  (Aldi, Walmart)")
bullet("Differentiation", "Highest perceived value → customers pay a premium  (Apple, Nike, Starbucks)")
bullet("Blue Ocean", "Both low cost AND high value simultaneously → uncontested new market space  (JetBlue)")
bullet("Commoditization", "Differentiation erodes when product becomes indistinguishable → premium pricing collapses")
bullet("Key insight", "Cost leader & differentiator CAN coexist in same industry (Aldi + Whole Foods in grocery)")

divider()

# ── Ch 8 ────────────────────────────────────────────────────────────────────
section_head("CH 8 — Corporate Strategy")
bullet("Corporate vs. Business", "Corporate = WHERE to compete (which industries).  Business = HOW to compete (cost leader vs. differentiator)")
bullet("Vertical integration", "Own more of the value chain → control, quality, proprietary knowledge  (Tesla: mfg + showrooms + charging)")
bullet("Outsourcing", "Contract non-core activities → focus on core competencies  (Nike: outsources mfg, owns design + brand)")
bullet("Diversification",
       "Same market + same competencies (easiest)  →  Related: one new  →  Unrelated: new market + new competencies (HARDEST)")
bullet("Economies of scale", "Spread fixed costs over more units → lower cost per unit  (Walmart buying power)")

divider()

# ── Ch 10 ───────────────────────────────────────────────────────────────────
section_head("CH 10 — Global Strategy")
bullet("CAGE Framework", "4 types of distance: Cultural · Administrative · Geographic · Economic  — higher distance = harder entry")
bullet("US → Canada", "Low CAGE distance (language, legal, border, income).   US → China/India = high CAGE distance")

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.15)
r = p.add_run("4 Global Strategy Types (Integration-Responsiveness 2×2):")
r.bold = True
r.font.size = Pt(9)
set_spacing(p, before=1, after=1)

bullet("Global",         "High integration, low local responsiveness — same product everywhere  (IKEA)", indent=0.35)
bullet("Multidomestic",  "Low integration, high local responsiveness — adapt to each market  (McDonald's local menus)", indent=0.35)
bullet("Transnational",  "High BOTH — hardest to execute; best of both worlds  (Unilever)", indent=0.35)
bullet("International",  "Low both — export home product with minimal changes  (early-stage exporters)", indent=0.35)

divider()

# ── Ch 11 ───────────────────────────────────────────────────────────────────
section_head("CH 11 — Organizational Design")
bullet("Why it matters",  "Org design is the EXECUTION layer — perfect strategy fails with the wrong structure/culture")
bullet("Culture",         "Shared values & norms = sustained competitive advantage because it cannot be copied  (Southwest Airlines)")
bullet("Inertia",         "Resistance to change embedded in routines, culture, incentives — past success breeds inertia  (Kodak had digital, couldn't adapt)")
bullet("Centralization",  "Top-down = control but slow.   Decentralization = fast but less control")
bullet("Business model",  "How strategy translates to $: value creation + delivery + capture  (Amazon flywheel)")

divider()

# ── Common Traps ─────────────────────────────────────────────────────────────
section_head("COMMON EXAM TRAPS — Do NOT fall for these")
trap("Differentiation ≠ lowest cost. Differentiators MANAGE cost but don't LEAD with it.")
trap("Cost leader ≠ cheapest price. Lowest cost STRUCTURE — price is a strategic choice.")
trap("Unrelated diversification ≠ a growth strategy. It's the hardest, riskiest move.")
trap("Vertical integration ≠ always superior. Sometimes outsourcing creates MORE competitive advantage.")
trap("Global strategy ≠ multidomestic. Global = standardize.  Multidomestic = adapt.")
trap("Culture ≠ 'company vibe.' It's a strategic asset that drives sustained competitive advantage.")
trap("Inertia ≠ laziness. It's structural — embedded in processes and incentives, not attitude.")

divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Every chapter answers one question: How do firms gain and sustain competitive advantage?  —  Good luck.")
r.italic = True
r.font.size = Pt(8.5)
set_spacing(p, before=2, after=0)

doc.save("StratMgmt_ExamDayOnePager.docx")
print("Saved: StratMgmt_ExamDayOnePager.docx")
